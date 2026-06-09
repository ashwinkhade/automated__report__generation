"""
Export Service.

Renders a generated report to HTML, PDF (via WeasyPrint) and DOCX (via python-docx).
HTML uses a self-contained Jinja2 template with embedded CSS so it can be opened
in any browser without external assets.
"""
from __future__ import annotations
import os
import logging
from typing import Dict, Any
from jinja2 import Template
from docx import Document
from docx.shared import Pt, RGBColor, Inches

logger = logging.getLogger(__name__)

HTML_TEMPLATE = """<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>{{ title }}</title>
<style>
  body { font-family: 'Segoe UI', Roboto, Arial, sans-serif; color:#1f2937; margin:40px; line-height:1.55; }
  h1 { color:#0f172a; border-bottom:3px solid #6366f1; padding-bottom:10px; }
  h2 { color:#4f46e5; margin-top:28px; }
  .meta { color:#64748b; font-size:13px; margin-bottom:24px; }
  .kpi-grid { display:grid; grid-template-columns:repeat(auto-fit, minmax(200px,1fr)); gap:14px; margin:18px 0; }
  .kpi { background:#f1f5f9; padding:14px 16px; border-radius:10px; border-left:4px solid #6366f1; }
  .kpi .label { font-size:12px; text-transform:uppercase; color:#64748b; letter-spacing:.05em; }
  .kpi .value { font-size:22px; font-weight:700; color:#0f172a; margin-top:4px; }
  ul { padding-left:22px; }
  li { margin-bottom:6px; }
  table { width:100%; border-collapse:collapse; margin-top:10px; font-size:13px; }
  th, td { border:1px solid #e2e8f0; padding:6px 10px; text-align:left; }
  th { background:#eef2ff; }
  .footer { margin-top:40px; padding-top:14px; border-top:1px solid #e2e8f0; color:#94a3b8; font-size:12px; }
</style>
</head>
<body>
  <h1>{{ title }}</h1>
  <div class="meta">Generated on {{ now }} · Automated Report Generation System</div>

  <h2>Executive Summary</h2>
  <p>{{ narrative.summary }}</p>

  <h2>Key Performance Indicators</h2>
  <div class="kpi-grid">
    {% for k, v in kpis.items() %}
      <div class="kpi">
        <div class="label">{{ k.replace('_',' ').title() }}</div>
        <div class="value">{{ v }}</div>
      </div>
    {% endfor %}
  </div>

  <h2>Key Insights</h2>
  <ul>
    {% for item in narrative.insights %}<li>{{ item }}</li>{% endfor %}
  </ul>

  <h2>Recommendations</h2>
  <ul>
    {% for item in narrative.recommendations %}<li>{{ item }}</li>{% endfor %}
  </ul>

  {% if charts %}
    <h2>Charts (Data Snapshot)</h2>
    {% for chart in charts %}
      <h3 style="color:#334155; margin-top:18px;">{{ chart.title }} ({{ chart.type }})</h3>
      <table>
        <thead><tr>{% for key in chart.data[0].keys() %}<th>{{ key }}</th>{% endfor %}</tr></thead>
        <tbody>
        {% for row in chart.data[:15] %}
          <tr>{% for v in row.values() %}<td>{{ v }}</td>{% endfor %}</tr>
        {% endfor %}
        </tbody>
      </table>
    {% endfor %}
  {% endif %}

  <div class="footer">© {{ year }} Automated Report Generation System — AI-powered analytics</div>
</body>
</html>
"""


class ExportService:
    """Render and save reports in multiple formats."""

    def to_html(self, path: str, title: str, narrative: Dict[str, Any], analytics: Dict[str, Any]) -> str:
        from datetime import datetime
        tpl = Template(HTML_TEMPLATE)
        html = tpl.render(
            title=title,
            narrative=narrative,
            kpis=analytics.get("kpis", {}),
            charts=analytics.get("charts", []),
            now=datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"),
            year=datetime.utcnow().year,
        )
        with open(path, "w", encoding="utf-8") as f:
            f.write(html)
        return path

    def to_pdf(self, path: str, title: str, narrative: Dict[str, Any], analytics: Dict[str, Any]) -> str:
        """Render HTML then convert to PDF using WeasyPrint."""
        html_str = self._render_html_string(title, narrative, analytics)
        try:
            from weasyprint import HTML
            HTML(string=html_str).write_pdf(path)
            return path
        except Exception as e:
            logger.warning("WeasyPrint failed (%s). Trying fallback.", e)
            # Fallback: write HTML with .pdf extension placeholder
            with open(path, "wb") as f:
                f.write(html_str.encode("utf-8"))
            return path

    def to_docx(self, path: str, title: str, narrative: Dict[str, Any], analytics: Dict[str, Any]) -> str:
        doc = Document()

        h = doc.add_heading(title, level=0)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

        from datetime import datetime
        doc.add_paragraph(
            f"Generated on {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}"
        ).italic = True

        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(narrative.get("summary", ""))

        doc.add_heading("Key Performance Indicators", level=1)
        kpis = analytics.get("kpis", {})
        if kpis:
            table = doc.add_table(rows=1, cols=2)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Metric"
            hdr[1].text = "Value"
            for k, v in kpis.items():
                row = table.add_row().cells
                row[0].text = k.replace("_", " ").title()
                row[1].text = str(v)

        doc.add_heading("Key Insights", level=1)
        for item in narrative.get("insights", []):
            doc.add_paragraph(item, style="List Bullet")

        doc.add_heading("Recommendations", level=1)
        for item in narrative.get("recommendations", []):
            doc.add_paragraph(item, style="List Bullet")

        doc.add_heading("Chart Data", level=1)
        for chart in analytics.get("charts", []):
            doc.add_heading(f"{chart['title']} ({chart['type']})", level=2)
            data = chart.get("data", [])
            if not data:
                continue
            keys = list(data[0].keys())
            table = doc.add_table(rows=1, cols=len(keys))
            table.style = "Light Grid Accent 1"
            for i, k in enumerate(keys):
                table.rows[0].cells[i].text = str(k)
            for row in data[:15]:
                cells = table.add_row().cells
                for i, k in enumerate(keys):
                    cells[i].text = str(row.get(k, ""))

        doc.save(path)
        return path

    def _render_html_string(self, title, narrative, analytics) -> str:
        from datetime import datetime
        tpl = Template(HTML_TEMPLATE)
        return tpl.render(
            title=title,
            narrative=narrative,
            kpis=analytics.get("kpis", {}),
            charts=analytics.get("charts", []),
            now=datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"),
            year=datetime.utcnow().year,
        )
